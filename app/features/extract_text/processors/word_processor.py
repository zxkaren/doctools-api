from pathlib import Path

from docx import Document


def extract_text_from_table(table) -> str:
    extracted_table_lines = []

    for table_row in table.rows:
        table_cells_text = []

        for table_cell in table_row.cells:
            cell_text = extract_text_from_cell(table_cell)

            if cell_text:
                table_cells_text.append(cell_text)

        if table_cells_text:
            extracted_table_lines.append(" | ".join(table_cells_text))

    return "\n".join(extracted_table_lines)


def extract_text_from_cell(table_cell) -> str:
    extracted_cell_blocks = []

    for paragraph in table_cell.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            extracted_cell_blocks.append(paragraph_text)

    # Recursividade para capturar textos em tabelas dentro de tabelas.
    for nested_table in table_cell.tables:
        nested_table_text = extract_text_from_table(nested_table)

        if nested_table_text:
            extracted_cell_blocks.append(nested_table_text)

    return " ".join(extracted_cell_blocks)


def extract_text_from_word(file_path: Path) -> str:
    """
    Resumo:
        Extrai apenas textos do corpo principal de um arquivo DOCX, incluindo
        textos em tabelas, sem capturar cabeçalhos, rodapés, imagens ou objetos.

    Parâmetros:
        file_path (Path): caminho do arquivo DOCX salvo em storage.

    Retorno:
        str: texto bruto extraído do documento.
    """
    extracted_text_blocks = []

    try:
        word_document = Document(file_path)

        for paragraph in word_document.paragraphs:
            paragraph_text = paragraph.text.strip()

            if paragraph_text:
                extracted_text_blocks.append(paragraph_text)

        for table in word_document.tables:
            table_text = extract_text_from_table(table)

            if table_text:
                extracted_text_blocks.append(table_text)

        return "\n\n".join(extracted_text_blocks)

    except Exception as error:
        print(f"Erro ao extrair texto do DOCX {file_path.name}: {error}")
        raise