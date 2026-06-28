import re
from copy import deepcopy
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ADD_RGB = RGBColor(0, 0, 255)
DELETE_RGB = RGBColor(255, 0, 0)

PARAGRAPH_BLOCK = "paragraph"
TABLE_BLOCK = "table"

MINIMUM_REPLACEMENT_SIMILARITY = 0.45


@dataclass
class DocumentBlock:
    block_type: str
    xml_element: Any
    content_object: Any
    comparison_text: str


def normalize_text_for_comparison(text: str) -> str:
    normalized_text = re.sub(r"\s+", " ", text or "")
    return normalized_text.strip().casefold()


def tokenize_with_separators(text: str) -> list[str]:
    return re.findall(r"\w+|[^\w\s]|\s+", text or "")


def count_significant_tokens(text: str) -> int:
    return sum(
        1
        for token in tokenize_with_separators(text)
        if not token.isspace()
    )


def build_token_run_pairs(paragraph: Paragraph) -> list[tuple[str, Any]]:
    token_run_pairs: list[tuple[str, Any]] = []

    for source_run in paragraph.runs:
        source_text = source_run.text or ""

        if not source_text:
            continue

        for token in tokenize_with_separators(source_text):
            token_run_pairs.append((token, source_run))

    return token_run_pairs


def get_significant_token_indices(token_run_pairs: list[tuple[str, Any]]) -> list[int]:
    significant_token_indices: list[int] = []

    for token_index, token_run_pair in enumerate(token_run_pairs):
        token_text, source_run = token_run_pair

        if source_run is not None and not token_text.isspace():
            significant_token_indices.append(token_index)

    return significant_token_indices


def copy_run_style(destination_run, source_run) -> None:
    try:
        destination_run.bold = source_run.bold
        destination_run.italic = source_run.italic
        destination_run.underline = source_run.underline
        destination_run.font.name = source_run.font.name
        destination_run.font.size = source_run.font.size

        if source_run.font.color and source_run.font.color.rgb is not None:
            destination_run.font.color.rgb = source_run.font.color.rgb

    except Exception:
        pass


def apply_change_style(run, change_type: str | None) -> None:
    if change_type == "add":
        run.font.color.rgb = ADD_RGB
        run.font.underline = True

    if change_type == "delete":
        run.font.color.rgb = DELETE_RGB
        run.font.strike = True


def write_tokens_range(
    paragraph: Paragraph,
    token_run_pairs: list[tuple[str, Any]],
    start_token_index: int,
    end_token_index: int,
    change_type: str | None = None,
) -> int:
    changed_tokens = 0

    for token_index in range(start_token_index, end_token_index):
        token_text, source_run = token_run_pairs[token_index]
        destination_run = paragraph.add_run(token_text)

        if source_run is not None:
            copy_run_style(destination_run, source_run)

        if change_type is not None:
            apply_change_style(destination_run, change_type)

            if not token_text.isspace():
                changed_tokens += 1

    return changed_tokens


def get_full_token_range(
    significant_token_indices: list[int],
    start_significant_index: int,
    end_significant_index: int,
    token_run_pairs_length: int,
) -> tuple[int, int]:
    if start_significant_index < len(significant_token_indices):
        start_token_index = significant_token_indices[start_significant_index]
    else:
        start_token_index = token_run_pairs_length

    if end_significant_index < len(significant_token_indices):
        end_token_index = significant_token_indices[end_significant_index]
    else:
        end_token_index = token_run_pairs_length

    return start_token_index, end_token_index


def update_paragraph_with_diff(
    compared_paragraph: Paragraph,
    original_paragraph: Paragraph,
    modified_paragraph: Paragraph,
) -> tuple[int, int]:
    """
    Compara dois parágrafos em nível de token e reescreve o parágrafo comparado
    preservando o máximo possível da formatação do arquivo modificado.

    Args:
        compared_paragraph: parágrafo que será reescrito no documento final.
        original_paragraph: parágrafo vindo do arquivo original.
        modified_paragraph: parágrafo vindo do arquivo modificado.

    Returns:
        Tupla com total de tokens adicionados e total de tokens removidos.
    """
    original_token_run_pairs = build_token_run_pairs(original_paragraph)
    modified_token_run_pairs = build_token_run_pairs(modified_paragraph)

    if not original_token_run_pairs and not modified_token_run_pairs:
        compared_paragraph.clear()
        return 0, 0

    original_significant_token_indices = get_significant_token_indices(original_token_run_pairs)
    modified_significant_token_indices = get_significant_token_indices(modified_token_run_pairs)

    original_significant_tokens = [
        normalize_text_for_comparison(original_token_run_pairs[token_index][0])
        for token_index in original_significant_token_indices
    ]

    modified_significant_tokens = [
        normalize_text_for_comparison(modified_token_run_pairs[token_index][0])
        for token_index in modified_significant_token_indices
    ]

    token_matcher = SequenceMatcher(
        None,
        original_significant_tokens,
        modified_significant_tokens,
        autojunk=False,
    )

    compared_paragraph.clear()

    total_additions = 0
    total_deletions = 0

    for opcode in token_matcher.get_opcodes():
        operation, original_start, original_end, modified_start, modified_end = opcode

        if operation == "equal":
            start_token_index, end_token_index = get_full_token_range(
                modified_significant_token_indices,
                modified_start,
                modified_end,
                len(modified_token_run_pairs),
            )

            write_tokens_range(
                compared_paragraph,
                modified_token_run_pairs,
                start_token_index,
                end_token_index,
            )

            continue

        if operation == "insert":
            start_token_index, end_token_index = get_full_token_range(
                modified_significant_token_indices,
                modified_start,
                modified_end,
                len(modified_token_run_pairs),
            )

            total_additions += write_tokens_range(
                compared_paragraph,
                modified_token_run_pairs,
                start_token_index,
                end_token_index,
                "add",
            )

            continue

        if operation == "delete":
            start_token_index, end_token_index = get_full_token_range(
                original_significant_token_indices,
                original_start,
                original_end,
                len(original_token_run_pairs),
            )

            total_deletions += write_tokens_range(
                compared_paragraph,
                original_token_run_pairs,
                start_token_index,
                end_token_index,
                "delete",
            )

            continue

        if operation == "replace":
            original_start_token_index, original_end_token_index = get_full_token_range(
                original_significant_token_indices,
                original_start,
                original_end,
                len(original_token_run_pairs),
            )

            modified_start_token_index, modified_end_token_index = get_full_token_range(
                modified_significant_token_indices,
                modified_start,
                modified_end,
                len(modified_token_run_pairs),
            )

            total_deletions += write_tokens_range(
                compared_paragraph,
                original_token_run_pairs,
                original_start_token_index,
                original_end_token_index,
                "delete",
            )

            total_additions += write_tokens_range(
                compared_paragraph,
                modified_token_run_pairs,
                modified_start_token_index,
                modified_end_token_index,
                "add",
            )

    return total_additions, total_deletions


def get_table_text(table: Table) -> str:
    row_texts: list[str] = []

    for table_row in table.rows:
        cell_texts = [table_cell.text for table_cell in table_row.cells]
        row_texts.append(" | ".join(cell_texts))

    return "\n".join(row_texts)


def collect_table_paragraphs_recursively(table: Table) -> list[Paragraph]:
    collected_paragraphs: list[Paragraph] = []

    for table_row in table.rows:
        for table_cell in table_row.cells:
            collected_paragraphs.extend(table_cell.paragraphs)

            for nested_table in table_cell.tables:
                collected_paragraphs.extend(
                    collect_table_paragraphs_recursively(nested_table)
                )

    return collected_paragraphs


def extract_body_blocks(document) -> list[DocumentBlock]:
    """
    Lê o corpo do DOCX na ordem real dos blocos XML.

    Essa abordagem evita comparar apenas `document.paragraphs` por índice,
    porque tabelas e parágrafos vazios podem alterar o fluxo visual do Word
    e desalinhar toda a comparação.

    Args:
        document: documento Word carregado pelo python-docx.

    Returns:
        Lista ordenada com blocos de parágrafo e tabela.
    """
    document_blocks: list[DocumentBlock] = []

    for xml_element in document.element.body.iterchildren():
        if xml_element.tag == qn("w:p"):
            paragraph = Paragraph(xml_element, document)
            comparison_text = normalize_text_for_comparison(paragraph.text)

            document_blocks.append(
                DocumentBlock(
                    block_type=PARAGRAPH_BLOCK,
                    xml_element=xml_element,
                    content_object=paragraph,
                    comparison_text=comparison_text,
                )
            )

            continue

        if xml_element.tag == qn("w:tbl"):
            table = Table(xml_element, document)
            comparison_text = normalize_text_for_comparison(get_table_text(table))

            document_blocks.append(
                DocumentBlock(
                    block_type=TABLE_BLOCK,
                    xml_element=xml_element,
                    content_object=table,
                    comparison_text=comparison_text,
                )
            )

    return document_blocks


def get_block_matcher_key(document_block: DocumentBlock) -> str:
    return f"{document_block.block_type}:{document_block.comparison_text}"


def clear_document_body(document) -> None:
    body_element = document.element.body
    section_properties = body_element.sectPr

    for child_element in list(body_element):
        if section_properties is not None and child_element is section_properties:
            continue

        body_element.remove(child_element)


def append_cloned_block(document, source_block: DocumentBlock) -> DocumentBlock:
    cloned_element = deepcopy(source_block.xml_element)
    body_element = document.element.body
    section_properties = body_element.sectPr

    if section_properties is not None:
        body_element.insert(body_element.index(section_properties), cloned_element)
    else:
        body_element.append(cloned_element)

    if source_block.block_type == PARAGRAPH_BLOCK:
        paragraph = Paragraph(cloned_element, document)

        return DocumentBlock(
            block_type=PARAGRAPH_BLOCK,
            xml_element=cloned_element,
            content_object=paragraph,
            comparison_text=source_block.comparison_text,
        )

    table = Table(cloned_element, document)

    return DocumentBlock(
        block_type=TABLE_BLOCK,
        xml_element=cloned_element,
        content_object=table,
        comparison_text=source_block.comparison_text,
    )


def apply_review_style_to_paragraph(paragraph: Paragraph, change_type: str) -> int:
    changed_tokens = 0

    for paragraph_run in paragraph.runs:
        run_text = paragraph_run.text or ""

        if not run_text:
            continue

        apply_change_style(paragraph_run, change_type)
        changed_tokens += count_significant_tokens(run_text)

    return changed_tokens


def apply_review_style_to_block(document_block: DocumentBlock, change_type: str) -> int:
    if document_block.block_type == PARAGRAPH_BLOCK:
        return apply_review_style_to_paragraph(document_block.content_object, change_type)

    table_text = get_table_text(document_block.content_object)

    if not normalize_text_for_comparison(table_text):
        return 0

    for table_paragraph in collect_table_paragraphs_recursively(document_block.content_object):
        apply_review_style_to_paragraph(table_paragraph, change_type)

    # Tabela inteira inserida/removida é tratada como uma alteração estrutural.
    return 1


def get_block_similarity(original_block: DocumentBlock, modified_block: DocumentBlock) -> float:
    if original_block.block_type != modified_block.block_type:
        return 0.0

    return SequenceMatcher(
        None,
        original_block.comparison_text,
        modified_block.comparison_text,
        autojunk=False,
    ).ratio()


def tables_have_same_structure(original_table: Table, modified_table: Table) -> bool:
    if len(original_table.rows) != len(modified_table.rows):
        return False

    for row_index, original_row in enumerate(original_table.rows):
        modified_row = modified_table.rows[row_index]

        if len(original_row.cells) != len(modified_row.cells):
            return False

    return True


def update_table_with_diff(
    compared_table: Table,
    original_table: Table,
    modified_table: Table,
) -> tuple[int, int]:
    total_additions = 0
    total_deletions = 0

    for row_index, compared_row in enumerate(compared_table.rows):
        original_row = original_table.rows[row_index]
        modified_row = modified_table.rows[row_index]

        for cell_index, compared_cell in enumerate(compared_row.cells):
            original_cell = original_row.cells[cell_index]
            modified_cell = modified_row.cells[cell_index]

            for paragraph_index, compared_paragraph in enumerate(compared_cell.paragraphs):
                if paragraph_index >= len(original_cell.paragraphs):
                    total_additions += apply_review_style_to_paragraph(
                        compared_paragraph,
                        "add",
                    )

                    continue

                original_paragraph = original_cell.paragraphs[paragraph_index]
                modified_paragraph = modified_cell.paragraphs[paragraph_index]

                added_count, deleted_count = update_paragraph_with_diff(
                    compared_paragraph,
                    original_paragraph,
                    modified_paragraph,
                )

                total_additions += added_count
                total_deletions += deleted_count

    return total_additions, total_deletions


def append_deleted_blocks(document, original_blocks: list[DocumentBlock]) -> int:
    total_deletions = 0

    for original_block in original_blocks:
        compared_block = append_cloned_block(document, original_block)
        total_deletions += apply_review_style_to_block(compared_block, "delete")

    return total_deletions


def append_added_blocks(document, modified_blocks: list[DocumentBlock]) -> int:
    total_additions = 0

    for modified_block in modified_blocks:
        compared_block = append_cloned_block(document, modified_block)
        total_additions += apply_review_style_to_block(compared_block, "add")

    return total_additions


def append_replaced_blocks(
    document,
    original_blocks: list[DocumentBlock],
    modified_blocks: list[DocumentBlock],
) -> tuple[int, int]:
    """
    Trata blocos marcados como `replace` pelo SequenceMatcher.

    Quando o bloco original e o modificado são parecidos, o código compara
    internamente em nível de token. Quando são muito diferentes, o bloco
    original é tratado como remoção e o bloco modificado como inserção.

    Args:
        document: documento comparado em construção.
        original_blocks: blocos substituídos vindos do documento original.
        modified_blocks: blocos substitutos vindos do documento modificado.

    Returns:
        Tupla com total de adições e total de exclusões.
    """
    total_additions = 0
    total_deletions = 0
    paired_blocks = min(len(original_blocks), len(modified_blocks))

    for block_index in range(paired_blocks):
        original_block = original_blocks[block_index]
        modified_block = modified_blocks[block_index]
        block_similarity = get_block_similarity(original_block, modified_block)

        if block_similarity < MINIMUM_REPLACEMENT_SIMILARITY:
            total_deletions += append_deleted_blocks(document, [original_block])
            total_additions += append_added_blocks(document, [modified_block])
            continue

        if original_block.block_type == TABLE_BLOCK:
            if not tables_have_same_structure(
                original_block.content_object,
                modified_block.content_object,
            ):
                total_deletions += append_deleted_blocks(document, [original_block])
                total_additions += append_added_blocks(document, [modified_block])
                continue

        compared_block = append_cloned_block(document, modified_block)

        if original_block.block_type == PARAGRAPH_BLOCK:
            added_count, deleted_count = update_paragraph_with_diff(
                compared_block.content_object,
                original_block.content_object,
                modified_block.content_object,
            )

            total_additions += added_count
            total_deletions += deleted_count

            continue

        if original_block.block_type == TABLE_BLOCK:
            added_count, deleted_count = update_table_with_diff(
                compared_block.content_object,
                original_block.content_object,
                modified_block.content_object,
            )

            total_additions += added_count
            total_deletions += deleted_count

    if len(original_blocks) > paired_blocks:
        total_deletions += append_deleted_blocks(
            document,
            original_blocks[paired_blocks:],
        )

    if len(modified_blocks) > paired_blocks:
        total_additions += append_added_blocks(
            document,
            modified_blocks[paired_blocks:],
        )

    return total_additions, total_deletions


def insert_summary_table(document, total_additions: int, total_deletions: int) -> None:
    summary_break_paragraph = document.add_paragraph()
    summary_break_run = summary_break_paragraph.add_run()
    summary_break_run.add_break(WD_BREAK.PAGE)

    summary_table = document.add_table(rows=3, cols=2)

    summary_table.rows[0].cells[0].text = "add"
    summary_table.rows[0].cells[1].text = str(total_additions)

    summary_table.rows[1].cells[0].text = "delete"
    summary_table.rows[1].cells[1].text = str(total_deletions)

    summary_table.rows[2].cells[0].text = "total_changes"
    summary_table.rows[2].cells[1].text = str(total_additions + total_deletions)


def compare_files(
    original_path: Path,
    modified_path: Path,
    processed_path: Path,
) -> dict[str, object]:
    """
    Compara dois arquivos DOCX e gera um documento final com marcações visuais.

    A comparação é feita em duas etapas:
    1. comparação estrutural por blocos do corpo do DOCX;
    2. comparação textual por tokens quando os blocos correspondentes são parecidos.

    Args:
        original_path: caminho do arquivo original.
        modified_path: caminho do arquivo modificado.
        processed_path: caminho onde o documento comparado será salvo.

    Returns:
        Dicionário com a tabela de resumo da comparação.
    """
    print("iniciando comparação de documentos Word")

    try:
        original_document = Document(original_path)
        modified_document = Document(modified_path)
        compared_document = Document(modified_path)

        original_blocks = extract_body_blocks(original_document)
        modified_blocks = extract_body_blocks(modified_document)

        original_matcher_keys = [
            get_block_matcher_key(document_block)
            for document_block in original_blocks
        ]

        modified_matcher_keys = [
            get_block_matcher_key(document_block)
            for document_block in modified_blocks
        ]

        block_matcher = SequenceMatcher(
            None,
            original_matcher_keys,
            modified_matcher_keys,
            autojunk=False,
        )

        clear_document_body(compared_document)

        total_additions = 0
        total_deletions = 0

        for opcode in block_matcher.get_opcodes():
            operation, original_start, original_end, modified_start, modified_end = opcode

            original_range_blocks = original_blocks[original_start:original_end]
            modified_range_blocks = modified_blocks[modified_start:modified_end]

            if operation == "equal":
                for modified_block in modified_range_blocks:
                    append_cloned_block(compared_document, modified_block)

                continue

            if operation == "insert":
                total_additions += append_added_blocks(
                    compared_document,
                    modified_range_blocks,
                )

                continue

            if operation == "delete":
                total_deletions += append_deleted_blocks(
                    compared_document,
                    original_range_blocks,
                )

                continue

            if operation == "replace":
                added_count, deleted_count = append_replaced_blocks(
                    compared_document,
                    original_range_blocks,
                    modified_range_blocks,
                )

                total_additions += added_count
                total_deletions += deleted_count

        insert_summary_table(
            compared_document,
            total_additions,
            total_deletions,
        )

        processed_path.parent.mkdir(parents=True, exist_ok=True)
        compared_document.save(processed_path)

        summary_table = {
            "add": total_additions,
            "delete": total_deletions,
            "total_changes": total_additions + total_deletions,
        }

        print("comparação de documentos Word concluída com sucesso")
        return {"summary_table": summary_table}

    except Exception as error_message:
        print(f"erro ao comparar documentos Word: {error_message}")
        raise